#!/usr/bin/env python3
"""
generate-hu-docx.py

Generates a .docx document in the exact format of the Metro de Medellin
"Anexo No 3 - Plantilla Modelamiento Historias de Usuario" template,
populated with approved user story data from data.json.

Strategy: "clear body, keep shell" -- opens the original template to
preserve headers (logo WMF), footers, page setup, and style definitions,
then clears all body content and rebuilds it programmatically.

Usage:
    python3 scripts/generate-hu-docx.py \
        --data output/RQ336/data.json \
        --template templates/core/metro-hu-base.docx \
        --output output/RQ336/MHU-RQ336.docx
"""

import sys
import os
import json
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Emu, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip3 install python-docx",
          file=sys.stderr)
    sys.exit(1)


# -- Constants (matching the original Metro template exactly) --------
FONT_NAME = 'Arial'
FONT_SIZE = Pt(10)
FONT_SIZE_9 = Pt(9)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_DARK_GRAY = RGBColor(0x20, 0x21, 0x24)
SHADING_GRAY = 'BFBFBF'


# -- Helpers ----------------------------------------------------------

def clear_body(doc):
    """Remove ALL content from the document body, keeping only sectPr
    (which holds page setup, margins, header/footer references).
    This also removes the TOC (sdt) and bookmarkEnd elements."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            body.remove(child)


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tc_pr.append(shading)


def add_bottom_border(paragraph, val='double', sz='4', space='1', color='auto'):
    """Add a bottom border to a paragraph (matching Heading 1 in template)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), val)
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_run(paragraph, text, bold=None, font_name=FONT_NAME,
            font_size=FONT_SIZE, color=COLOR_BLACK, italic=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    return run


def add_para(doc, text='', alignment=None, bold=None, italic=None,
             font_size=FONT_SIZE, color=COLOR_BLACK,
             space_before=None, space_after=None):
    """Add a simple paragraph with a single run."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    else:
        pf.space_after = Pt(0)
    if text:
        add_run(p, text, bold=bold, font_size=font_size,
                color=color, italic=italic)
    return p


def add_heading1(doc, text):
    """Add a Heading 1 paragraph matching the original template:
    Style=Heading 1, but run overrides to Arial 10pt black uppercase,
    with a double bottom border."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['Heading 1']
    except KeyError:
        pass
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(24)
    pf.space_after = Pt(0)
    add_run(p, text.upper(), font_name=FONT_NAME, font_size=FONT_SIZE,
            color=COLOR_BLACK)
    add_bottom_border(p)
    return p


def add_heading2(doc, text):
    """Add a Heading 2 paragraph matching the original template:
    Style=Heading 2, but run overrides to Arial 10pt black."""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles['Heading 2']
    except KeyError:
        pass
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(10)
    pf.space_after = Pt(0)
    add_run(p, text, font_name=FONT_NAME, font_size=FONT_SIZE,
            color=COLOR_BLACK)
    return p


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_table_with_header(doc, headers, data_rows):
    """Create a table with a gray header row and data rows."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.style = 'Table Grid'
    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, SHADING_GRAY)
        p = cell.paragraphs[0]
        add_run(p, header_text, font_name=FONT_NAME,
                font_size=FONT_SIZE, color=COLOR_BLACK)
    # Data rows
    for r, row_data in enumerate(data_rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            p = cell.paragraphs[0]
            add_run(p, str(cell_text), font_name=FONT_NAME,
                    font_size=FONT_SIZE_9, color=COLOR_BLACK)
    return table


# -- Section builders -------------------------------------------------

def build_cover_page(doc, sprint_id, proyecto_nombre, fecha_str):
    """Cover page: centered, TP-like formatting. Separate page."""
    for _ in range(5):
        add_para(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(0))
    add_para(doc, 'MHU', alignment=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=FONT_SIZE, color=COLOR_BLACK, italic=False,
             space_before=Pt(0), space_after=Pt(0))
    add_para(doc, sprint_id, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=FONT_SIZE, color=COLOR_BLACK, italic=False,
             space_before=Pt(0), space_after=Pt(0))
    for _ in range(2):
        add_para(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=Pt(0), space_after=Pt(0))
    add_para(doc, f'({proyecto_nombre})', alignment=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=FONT_SIZE, color=COLOR_BLACK, italic=False,
             space_before=Pt(0), space_after=Pt(0))
    add_para(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # "Version 1.0" bold centered
    p_ver = add_para(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p_ver, 'Version ', bold=True, font_size=FONT_SIZE, color=COLOR_BLACK)
    add_run(p_ver, '1.0', bold=True, font_size=FONT_SIZE, color=COLOR_BLACK)
    # Date
    add_para(doc, fecha_str, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             font_size=FONT_SIZE, color=COLOR_BLACK)
    for _ in range(6):
        add_para(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # PAGE BREAK after cover
    add_page_break(doc)


def build_control_tables(doc, fecha_str, sprint_id):
    """Control de Versiones, Revisiones, Aprobaciones. Separate page."""
    # Control de Versiones
    add_para(doc, 'Control de Versiones', bold=True)
    add_para(doc, '')
    add_table_with_header(
        doc,
        headers=['Fecha', 'Version', 'Descripcion', 'Autor'],
        data_rows=[
            [fecha_str, '1.0', sprint_id, ''],
            ['', '', '', ''],
            ['', '', '', ''],
        ]
    )
    add_para(doc, '')
    # Revisiones
    add_para(doc, 'Revisiones', bold=True)
    add_para(doc, '')
    add_table_with_header(
        doc,
        headers=['Revisor', 'Cargo', 'Fecha'],
        data_rows=[['', '', ''], ['', '', ''], ['', '', '']]
    )
    add_para(doc, '')
    # Aprobaciones
    add_para(doc, 'Aprobaciones', bold=True)
    add_para(doc, '')
    add_table_with_header(
        doc,
        headers=['Revisor', 'Cargo', 'Fecha'],
        data_rows=[['', '', ''], ['', '', ''], ['', '', '']]
    )
    # PAGE BREAK after control tables
    add_page_break(doc)


def build_introduction(doc, data, fecha_str):
    """Section 1: INTRODUCCION with subsections 1.1-1.6."""
    historias = data.get('historias', [])
    config = data.get('config', {})
    sprint_config = data.get('sprint_config', {})
    proyecto = config.get('proyecto_nombre', 'Proyecto')
    sprint_id = data.get('meta', {}).get('sprint_id', 'Sprint-X')
    total_horas = sum(h.get('estimacion_total_horas', 0) for h in historias)
    fecha_fin = sprint_config.get('fecha_fin', 'por definir')
    fecha_inicio = sprint_config.get('fecha_inicio', fecha_str)

    # Derive business-oriented descriptions from HU data
    hu_titles = [h.get('titulo_mejorado') or h.get('titulo_original') or h['hu_id']
                 for h in historias]
    hu_benefits = [h.get('narrativa_refinada', {}).get('beneficio', '')
                   for h in historias if h.get('narrativa_refinada', {}).get('beneficio')]

    capabilities_summary = '; '.join(hu_titles[:5])
    business_value = '. '.join(
        [b for b in hu_benefits if b and len(b) > 10][:3]
    ) if hu_benefits else 'optimizar los procesos operativos del negocio'

    # --- Heading ---
    add_heading1(doc, 'INTRODUCCION')
    add_para(doc, '')
    add_para(doc,
             f'El presente documento describe el modelamiento de las historias '
             f'de usuario correspondientes al requerimiento {sprint_id} del '
             f'proyecto {proyecto}.')
    add_para(doc, '')
    add_para(doc,
             f'Se definen {len(historias)} historias de usuario que cubren '
             f'las siguientes capacidades: {capabilities_summary}.')

    # --- 1.1 Glosario ---
    add_para(doc, '')
    add_heading2(doc, 'Glosario')
    add_para(doc, '')

    # Build glossary from domain concepts in HU data
    glossary = _build_glossary(historias, proyecto)
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f'\u2022 {term}: ', bold=True, font_size=FONT_SIZE, color=COLOR_BLACK)
        add_run(p, definition, font_size=FONT_SIZE, color=COLOR_BLACK)

    # --- 1.2 Objetivo del Negocio ---
    add_para(doc, '')
    add_heading2(doc, 'Objetivo del Negocio')
    add_para(doc, '')
    add_para(doc, business_value if len(business_value) > 30 else
             f'Habilitar las capacidades operativas y funcionales requeridas '
             f'por {proyecto} para {business_value}.')

    # --- 1.3 Objetivo SMART ---
    add_para(doc, '')
    add_heading2(doc, 'Objetivo SMART')

    smart = _build_smart(historias, total_horas, fecha_fin, proyecto)
    for label, value in smart:
        add_para(doc, '')
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f'\u00bF{label} ', bold=True, color=COLOR_DARK_GRAY)
        add_run(p, value, font_size=FONT_SIZE, color=COLOR_DARK_GRAY)

    # --- 1.4 Requisitos fuera del alcance ---
    add_para(doc, '')
    add_heading2(doc, 'Requisitos fuera del alcance')
    add_para(doc, '')

    # --- 1.5 Riesgos Generales ---
    add_para(doc, '')
    add_heading2(doc, 'Riesgos Generales')
    add_para(doc, '')

    all_risks = []
    for h in historias:
        for r in h.get('riesgos', []):
            nivel = r.get('nivel_riesgo_valor',
                          r.get('gravedad', 1) * r.get('probabilidad_riesgo', 1))
            if int(nivel) >= 10:
                all_risks.append(r)

    if all_risks:
        for r in all_risks[:10]:
            desc = r.get('descripcion', '')
            tipologia = r.get('tipologia', r.get('categoria', ''))
            nivel_label = r.get('nivel_riesgo_label', 'Alto')
            add_para(doc,
                     f'\u2022 [{r.get("riesgo_id", "")}] {tipologia}: '
                     f'{desc} ({nivel_label})')
    else:
        add_para(doc, 'Sin riesgos criticos identificados.')

    # --- 1.6 Referencias ---
    add_para(doc, '')
    add_heading2(doc, 'Referencias')
    add_para(doc, '')


def _build_glossary(historias, proyecto):
    """Extract key domain terms from HU data for a meaningful glossary."""
    terms = []
    # Fixed terms
    terms.append(('HU', 'Historia de Usuario. Unidad funcional de entrega.'))
    terms.append(('CA', 'Criterio de Aceptacion. Condicion verificable que '
                        'una HU debe cumplir para considerarse completa.'))
    terms.append(('RF', 'Requerimiento Funcional. Agrupacion logica de '
                        'historias de usuario relacionadas.'))

    # Extract domain-specific terms from HU titles and roles
    seen = set()
    for h in historias:
        nr = h.get('narrativa_refinada', {})
        rol = nr.get('rol', '')
        if rol and rol not in seen and len(rol) > 5:
            seen.add(rol)
            terms.append((rol.split('(')[0].strip(),
                          f'Actor del sistema que interactua con la '
                          f'funcionalidad descrita en {h.get("hu_id", "la HU")}.'))
            if len(terms) >= 8:
                break

    return terms


def _build_smart(historias, total_horas, fecha_fin, proyecto):
    """Build SMART objective items with meaningful, specific content."""
    # Que: what the stories actually declare
    hu_summaries = []
    for h in historias:
        titulo = h.get('titulo_mejorado') or h.get('titulo_original') or h['hu_id']
        hu_summaries.append(titulo)
    que_text = '; '.join(hu_summaries)
    if len(que_text) > 300:
        que_text = que_text[:297] + '...'

    # Cuanto: effort
    cuanto_text = f'{total_horas:.1f} horas estimadas de desarrollo.'

    # Como: implementation strategy from tasks
    task_categories = set()
    for h in historias:
        for t in h.get('tareas', []):
            titulo = t.get('titulo', t.get('descripcion', ''))
            if titulo:
                # Extract first verb/concept
                first_word = titulo.split()[0] if titulo.split() else ''
                if len(first_word) > 3:
                    task_categories.add(first_word)
    if task_categories:
        top_tasks = list(task_categories)[:5]
        como_text = (f'Desarrollo iterativo abarcando: '
                     f'{", ".join(top_tasks).lower()}.')
    else:
        como_text = 'Desarrollo iterativo con validacion continua de criterios de aceptacion.'

    # Para que: specific business value from benefits
    benefits = []
    for h in historias:
        b = h.get('narrativa_refinada', {}).get('beneficio', '')
        if b and len(b) > 15:
            benefits.append(b)
    if benefits:
        para_que_text = benefits[0]
        if len(benefits) > 1:
            para_que_text += f'; y adicionalmente, {benefits[1].lower()}'
        if len(para_que_text) > 300:
            para_que_text = para_que_text[:297] + '...'
    else:
        para_que_text = f'Entregar las capacidades funcionales definidas para {proyecto}.'

    # Cuando: sprint end date
    cuando_text = fecha_fin

    return [
        ('Que?', que_text),
        ('Cuanto?', cuanto_text),
        ('Como?', como_text),
        ('Para que?', para_que_text),
        ('Cuando?', cuando_text),
    ]


def build_hu_rf_mapping(doc, historias):
    """Section 2: RELACION DE HU / HABILITADORES VS RF."""
    add_heading1(doc, 'RELACION DE HU / HABILITADORES VS RF')
    add_para(doc, '')
    add_para(doc,
             'Relacion de las historias de usuario o habilitadores '
             'vs los requisitos funcionales / no funcionales')
    add_para(doc, '')
    data_rows = []
    for idx, h in enumerate(historias):
        rf_id = f'RF{idx + 1:02d}'
        hu_id = h.get('hu_id', f'HU{idx + 1:02d}')
        data_rows.append([rf_id, hu_id])
    add_table_with_header(
        doc,
        headers=['Requerimientos Funcionales\n/ No Funcionales',
                 'Historias de Usuario / Habilitadores'],
        data_rows=data_rows
    )


def build_hu_specifications(doc, historias):
    """Section 3: ESPECIFICACION DE HISTORIAS DE USUARIO."""
    add_heading1(doc, 'ESPECIFICACION DE HISTORIAS DE USUARIO')
    add_para(doc, '')

    for idx, h in enumerate(historias):
        rf_id = f'RF{idx + 1:02d}'
        hu_id = h.get('hu_id', f'HU{idx + 1:02d}')
        nr = h.get('narrativa_refinada', {})
        rol = nr.get('rol', 'usuario')
        accion = nr.get('accion', '(sin especificar)')
        beneficio = nr.get('beneficio', '(sin especificar)')
        criterios = h.get('criteriosAceptacion', [])

        # RF heading (Heading 2 = Arial 10pt black)
        add_heading2(doc, rf_id)
        add_para(doc, '')

        # HU ID (bold, indented)
        p_hu = doc.add_paragraph()
        p_hu.paragraph_format.left_indent = Inches(0.5)
        p_hu.paragraph_format.line_spacing = 1.15
        p_hu.paragraph_format.space_after = Pt(0)
        add_run(p_hu, f'{hu_id} ', bold=True, font_size=FONT_SIZE,
                color=COLOR_BLACK)

        add_para(doc, '')

        # "Yo como {rol}"
        p_yo = doc.add_paragraph()
        p_yo.paragraph_format.line_spacing = 1.15
        p_yo.paragraph_format.space_after = Pt(0)
        add_run(p_yo, 'Yo', bold=True, font_size=FONT_SIZE)
        add_run(p_yo, f' como {rol}', font_size=FONT_SIZE)

        # "Deseo {accion}"
        p_deseo = doc.add_paragraph()
        p_deseo.paragraph_format.line_spacing = 1.15
        p_deseo.paragraph_format.space_after = Pt(0)
        add_run(p_deseo, 'Deseo', bold=True, font_size=FONT_SIZE)
        add_run(p_deseo, f' {accion}', font_size=FONT_SIZE)

        # "Para {beneficio}"
        p_para = doc.add_paragraph()
        p_para.paragraph_format.line_spacing = 1.15
        p_para.paragraph_format.space_after = Pt(0)
        add_run(p_para, 'Para', bold=True, font_size=FONT_SIZE)
        add_run(p_para, f' {beneficio}', font_size=FONT_SIZE)

        add_para(doc, '')

        # "Criterios de Aceptacion"
        p_ca = doc.add_paragraph()
        p_ca.paragraph_format.line_spacing = 1.15
        p_ca.paragraph_format.space_after = Pt(0)
        add_run(p_ca, 'Criterios de Aceptacion', bold=True,
                font_size=FONT_SIZE)

        # Numbered criteria
        for ci, ca in enumerate(criterios):
            nombre = (ca.get('nombre') or ca.get('escenario_id')
                      or f'Criterio {ci + 1}')
            add_para(doc, f'{ci + 1}.        {nombre}')

        add_para(doc, '')


def build_habilitadores(doc):
    """Section 4: ESPECIFICACION DE HABILITADORES."""
    add_heading1(doc, 'ESPECIFICACION DE HABILITADORES')
    add_para(doc, '')
    add_heading2(doc, 'RNF')
    add_para(doc, '')
    p_h = doc.add_paragraph()
    p_h.paragraph_format.left_indent = Inches(0.5)
    p_h.paragraph_format.line_spacing = 1.15
    add_run(p_h, 'Ver seccion de riesgos y dependencias en cada '
                 'historia de usuario.',
            font_size=FONT_SIZE, color=COLOR_BLACK)


# -- Main -------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Generate .docx from data.json using Metro template')
    parser.add_argument('--data', required=True, help='Path to data.json')
    parser.add_argument('--template', required=True,
                        help='Path to metro-hu-base.docx')
    parser.add_argument('--output', required=True, help='Output .docx path')
    args = parser.parse_args()

    if not os.path.exists(args.data):
        print(f'ERROR: data.json not found: {args.data}', file=sys.stderr)
        sys.exit(1)
    if not os.path.exists(args.template):
        print(f'ERROR: template not found: {args.template}', file=sys.stderr)
        sys.exit(1)

    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)

    historias = data.get('historias', [])
    config = data.get('config', {})
    meta = data.get('meta', {})
    sprint_id = meta.get('sprint_id', 'Sprint-X')
    proyecto_nombre = config.get('proyecto_nombre', 'Proyecto')
    fecha_str = datetime.now().strftime('%d-%m-%Y')

    # Open template (preserves header, footer, styles, logo)
    doc = Document(args.template)
    clear_body(doc)

    # Build document in exact order:
    # Page 1: Cover page
    build_cover_page(doc, sprint_id, proyecto_nombre, fecha_str)
    # Page 2: Control tables
    build_control_tables(doc, fecha_str, sprint_id)
    # Page 3+: Content
    build_introduction(doc, data, fecha_str)
    build_hu_rf_mapping(doc, historias)
    build_hu_specifications(doc, historias)
    build_habilitadores(doc)

    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
    doc.save(args.output)
    size_kb = os.path.getsize(args.output) / 1024
    print(f'OK: {args.output} ({size_kb:.1f} KB, {len(historias)} HUs)')


if __name__ == '__main__':
    main()
